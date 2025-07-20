import streamlit as st
import streamlit.components.v1 as components
from docx import Document
import re
import uuid
import os
import time
import html
import csv
from io import BytesIO

st.set_page_config(
    page_title="Ø§Ù„Ù‚ÙˆØ§Ù†ÙŠÙ† Ø§Ù„ÙŠÙ…Ù†ÙŠØ© Ø¨Ø¢Ø®Ø± ØªØ¹Ø¯ÙŠÙ„Ø§ØªÙ‡Ø§ Ø­ØªÙ‰ Ø¹Ø§Ù… 2025Ù…",
    layout="wide",
    initial_sidebar_state="collapsed"
)

if "night_mode" not in st.session_state:
    st.session_state.night_mode = False

st.markdown("""
<style>
textarea, input[type="text"], .stTextArea textarea, .stTextInput input {
    direction: rtl !important;
    text-align: right !important;
    font-family: "Tahoma", "Arial", sans-serif !important;
    font-size: 18px !important;
}
[data-testid="stTextArea"] textarea,
[data-testid="stTextInput"] input {
    direction: rtl !important;
    text-align: right !important;
}
mark {
    background: #ff9800 !important;
    color: #fff !important;
}
mark.mark-soft {
    background: #ffd600 !important;
    color: #000 !important;
}
</style>
""", unsafe_allow_html=True)

TRIAL_DURATION = 3 * 24 * 60 * 60
TRIAL_USERS_FILE = "trial_users.txt"
DEVICE_ID_FILE = "device_id.txt"
ACTIVATED_FILE = "activated.txt"
ACTIVATION_CODES_FILE = "activation_codes.txt"
LAWS_DIR = "laws"

def get_device_id():
    if os.path.exists(DEVICE_ID_FILE):
        with open(DEVICE_ID_FILE, "r") as f:
            return f.read().strip()
    new_id = str(uuid.uuid4())
    with open(DEVICE_ID_FILE, "w") as f:
        f.write(new_id)
    return new_id

def get_trial_start(device_id):
    if not os.path.exists(TRIAL_USERS_FILE):
        return None
    with open(TRIAL_USERS_FILE, "r") as f:
        reader = csv.reader(f)
        for row in reader:
            if row and row[0] == device_id:
                return float(row[1])
    return None

def register_trial(device_id):
    if not os.path.exists(TRIAL_USERS_FILE):
        with open(TRIAL_USERS_FILE, "w", newline='') as f:
            pass
    with open(TRIAL_USERS_FILE, "a", newline='') as f:
        writer = csv.writer(f)
        writer.writerow([device_id, time.time()])

def is_activated():
    return os.path.exists(ACTIVATED_FILE)

def activate_app(code):
    if not os.path.exists(ACTIVATION_CODES_FILE):
        return False
    with open(ACTIVATION_CODES_FILE, "r") as f:
        codes = [line.strip() for line in f.readlines()]
    if code in codes:
        codes.remove(code)
        with open(ACTIVATION_CODES_FILE, "w") as f:
            for c in codes:
                f.write(c + "\n")
        with open(ACTIVATED_FILE, "w") as f:
            f.write("activated")
        return True
    return False

def highlight_keywords(text, keywords, normalized_keywords=None, exact_match=False):
    if not keywords:
        return text
    marked_spans = []
    for kw in keywords:
        if not kw:
            continue
        for m in re.finditer(r'(?<!\w)' + re.escape(kw) + r'(?!\w)', text, re.IGNORECASE):
            marked_spans.append((m.start(), m.end(), "exact"))
    if normalized_keywords:
        normalized_text = normalize_arabic_text(text)
        for i, norm_kw in enumerate(normalized_keywords):
            if not norm_kw:
                continue
            original_kw = keywords[i]
            if not exact_match:
                for m in re.finditer(re.escape(original_kw), text, re.IGNORECASE):
                    overlap = False
                    for s, e, t in marked_spans:
                        if not (m.end() <= s or m.start() >= e):
                            overlap = True
                            break
                    if not overlap:
                        marked_spans.append((m.start(), m.end(), "partial"))
    if not marked_spans:
        return text
    marked_spans.sort(key=lambda x: x[0])
    result = []
    last_idx = 0
    for s, e, t in marked_spans:
        if s < last_idx:
            continue
        result.append(text[last_idx:s])
        span_text = text[s:e]
        if t == "exact":
            result.append(f"<mark>{span_text}</mark>")
        else:
            result.append(f"<mark class=\"mark-soft\">{span_text}</mark>")
        last_idx = e
    result.append(text[last_idx:])
    return "".join(result)

def export_results_to_word(results, filename="Ù†ØªØ§Ø¦Ø¬_Ø§Ù„Ø¨Ø­Ø«.docx"):
    document = Document()
    document.add_heading('Ù†ØªØ§Ø¦Ø¬ Ø§Ù„Ø¨Ø­Ø« ÙÙŠ Ø§Ù„Ù‚ÙˆØ§Ù†ÙŠÙ† Ø§Ù„ÙŠÙ…Ù†ÙŠØ©', level=1)
    if not results:
        document.add_paragraph("Ù„Ù… ÙŠØªÙ… Ø§Ù„Ø¹Ø«ÙˆØ± Ø¹Ù„Ù‰ Ù†ØªØ§Ø¦Ø¬ Ù„Ù„ÙƒÙ„Ù…Ø§Øª Ø§Ù„Ù…ÙØªØ§Ø­ÙŠØ© Ø§Ù„Ù…Ø­Ø¯Ø¯Ø©.")
    else:
        for i, r in enumerate(results):
            document.add_heading(f"Ø§Ù„Ù‚Ø§Ù†ÙˆÙ†: {r['law']} - Ø§Ù„Ù…Ø§Ø¯Ø©: {r['num']}", level=2)
            document.add_paragraph(r['plain'])
            if i < len(results) - 1:
                document.add_page_break()
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def normalize_arabic_numbers(text):
    arabic_to_english = str.maketrans('Ù Ù¡Ù¢Ù£Ù¤Ù¥Ù¦Ù§Ù¨Ù©', '0123456789')
    return text.translate(arabic_to_english)

def normalize_arabic_text(text):
    text = re.sub(r'(.)\1{2,}', r'\1', text)
    text = re.sub(r'[\u064B-\u0652]', '', text)
    text = re.sub('[Ø¥Ø£Ø¢Ø§]', 'Ø§', text)
    text = re.sub('[Ù‰ÙŠ]', 'ÙŠ', text)
    text = re.sub('[Ø©]', 'Ù‡', text)
    text = re.sub('Ø¤', 'Ùˆ', text)
    text = re.sub('Ø¦', 'ÙŠ', text)
    text = re.sub(r'[^\w\s]', '', text)
    text = re.sub('\s+', ' ', text)
    return text.strip()

def render_law_file_viewer(files):
    st.markdown("<h4 style='text-align:center;'>Ø§Ø®ØªØ± Ø§Ù„Ù‚Ø§Ù†ÙˆÙ† Ø§Ù„Ø°ÙŠ ØªØ±ÙŠØ¯ ØªØµÙØ­Ù‡ Ø¨Ø§Ù„ÙƒØ§Ù…Ù„:</h4>", unsafe_allow_html=True)
    law_sel = st.selectbox("Ø§Ø®ØªØ± Ø§Ù„Ù‚Ø§Ù†ÙˆÙ†:", files, key="law_select_for_view")
    if law_sel:
        doc = Document(os.path.join(LAWS_DIR, law_sel))
        st.markdown(f"<h5 style='text-align:center;color:#1976d2'>{law_sel.replace('.docx','')}</h5>", unsafe_allow_html=True)
        law_text = ""
        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt:
                law_text += txt + "\n\n"
        st.markdown("""
        <style>
        textarea[disabled], .stTextArea textarea[disabled] {
            color: #000 !important;
            background: #fff !important;
            opacity: 1 !important;
            filter: none !important;
            font-size: 19px !important;
            font-family: "Tahoma", "Arial", sans-serif !important;
            font-weight: bold !important;
            letter-spacing: 0.3px;
        }
        textarea::-webkit-textfield-decoration-container {
            display: none !important;
        }
        textarea::-webkit-scrollbar-button,
        textarea::-webkit-scrollbar-corner {
            display: none !important;
        }
        textarea::selection { background: #b3d7ff; }
        textarea[readonly]::-moz-selection,
        textarea[disabled]::-moz-selection {
            background: #b3d7ff;
        }
        textarea[disabled]::selection {
            background: #b3d7ff;
        }
        </style>
        """, unsafe_allow_html=True)
        st.text_area("Ø§Ù„Ù‚Ø§Ù†ÙˆÙ† ÙƒØ§Ù…Ù„:", law_text, height=550, key="full_law_view_text", disabled=True)

def run_main_app():
    with st.sidebar:
        col1, col2 = st.columns([1, 1])
        with col1:
            st.session_state.night_mode = st.toggle("ğŸŒ™ ØªÙØ¹ÙŠÙ„ Ø§Ù„ÙˆØ¶Ø¹ Ø§Ù„Ù„ÙŠÙ„ÙŠ", value=st.session_state.night_mode)
        with col2:
            whatsapp_number = "+967777533034"
            whatsapp_msg = "Ù…Ø±Ø­Ø¨Ù‹Ø§ØŒ Ù„Ø¯ÙŠ Ø§Ø³ØªÙØ³Ø§Ø± Ø¨Ø®ØµÙˆØµ ØªØ·Ø¨ÙŠÙ‚ Ø§Ù„Ù‚ÙˆØ§Ù†ÙŠÙ† Ø§Ù„ÙŠÙ…Ù†ÙŠØ©."
            whatsapp_url = f"https://wa.me/{whatsapp_number}?text={whatsapp_msg.replace(' ', '%20')}"
            st.markdown(
                f"""
                <a href="{whatsapp_url}" target="_blank" style="display:inline-block;text-align:center;">
                    <img src="https://img.icons8.com/color/48/000000/whatsapp--v1.png" width="32" style="vertical-align:middle;margin-bottom:6px;" alt="ÙˆØ§ØªØ³Ø§Ø¨"/>
                    <br style="display:none"/>
                    <span style="font-size:12px;">Ù…Ø±Ø§Ø³Ù„ØªÙ†Ø§</span>
                </a>
                """,
                unsafe_allow_html=True,
            )
    tabs = st.tabs(["ğŸ” Ø§Ù„Ø¨Ø­Ø« ÙÙŠ Ø§Ù„Ù‚ÙˆØ§Ù†ÙŠÙ†", "ğŸ“„ Ø¹Ø±Ø¶ Ø§Ù„Ù‚Ø§Ù†ÙˆÙ† Ø§Ù„ÙƒØ§Ù…Ù„"])
    with tabs[0]:
        if st.session_state.night_mode:
            st.markdown("""
            <style>
            body, .stApp {
                background-color: #181a1b !important;
                color: #f1f1f1 !important;
            }
            .stTextInput input, .stTextArea textarea, textarea, input[type="text"] {
                background-color: #222426 !important;
                color: #f1f1f1 !important;
            }
            .stButton button, .stDownloadButton button {
                background: linear-gradient(90deg, #333 0%, #222 100%) !important;
                color: #f1f1f1 !important;
                border: 1px solid #444 !important;
            }
            .stExpanderHeader, .stForm, .stMetric {
                background-color: #232526 !important;
                color: #f1f1f1 !important;
            }
            mark {
                background: #ff9800 !important;
                color: #fff !important;
            }
            mark.mark-soft {
                background: #ffd600 !important;
                color: #000 !important;
            }
            .copy-material-btn {
                background: linear-gradient(90deg, #384e5a 0%, #213b4b 100%) !important;
                color: #eee !important;
            }
            .copy-material-btn:hover {
                background: linear-gradient(90deg, #213b4b 0%, #384e5a 100%) !important;
            }
            .result-box-night {
                background-color: #232526 !important;
                color: #fafafa !important;
                padding: 20px;
                margin-bottom: 10px;
                width: 100%;
                max-width: 100%;
                border-radius: 10px;
                border: 1px solid #333;
                direction: rtl;
                text-align: right;
            }
            </style>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <style>
            body, .stApp {
                background-color: #fff !important;
                color: #232323 !important;
            }
            mark {
                background: #ff9800 !important;
                color: #fff !important;
            }
            mark.mark-soft {
                background: #ffd600 !important;
                color: #000 !important;
            }
            .result-box-night {
                background-color: #f1f8e9 !important;
                color: #232323 !important;
                padding: 20px;
                margin-bottom: 10px;
                width: 100%;
                max-width: 100%;
                border-radius: 10px;
                border: 1px solid #c5e1a5;
                direction: rtl;
                text-align: right;
            }
            </style>
            """, unsafe_allow_html=True)
        components.html("""
        <style>
        .scroll-btn {
            position: fixed;
            left: 10px;
            padding: 12px;
            font-size: 24px;
            border-radius: 50%;
            background-color: #c5e1a5;
            color: black;
            cursor: pointer;
            z-index: 9999;
            border: none;
            box-shadow: 1px 1px 5px #888;
        }
        #scroll-top-btn { bottom: 80px; }
        #scroll-bottom-btn { bottom: 20px; }
        .rtl-metric {
            direction: rtl;
            text-align: right !important;
            margin-right: 0 !important;
        }
        .rtl-metric .stMetric {
            text-align: right !important;
            direction: rtl;
        }
        .rtl-metric .stMetricDelta {
            display: block !important;
            text-align: right !important;
            direction: rtl;
        }
        .rtl-download-btn {
            direction: rtl;
            text-align: right !important;
            margin-right: 0 !important;
            display: flex;
            flex-direction: row-reverse;
            justify-content: flex-start;
        }
        textarea, .stTextArea, .stTextArea textarea, input[type="text"], .stTextInput input, .stTextInput textarea {
            direction: rtl !important;
            text-align: right !important;
            padding-right: 10px;
            font-family: "Tahoma", "Arial", sans-serif;
            font-size: 16px;
            line-height: 1.5;
        }
        .stButton, .stDownloadButton, .stMetric {
            direction: rtl !important;
            text-align: right !important;
        }
        </style>
        <button class='scroll-btn' id='scroll-top-btn' onclick='window.scrollTo({top: 0, behavior: "smooth"});'>â¬†ï¸</button>
        <button class='scroll-btn' id='scroll-bottom-btn' onclick='window.scrollTo({top: document.body.scrollHeight, behavior: "smooth"});'>â¬‡ï¸</button>
        """, height=1)
        if not os.path.exists(LAWS_DIR):
            st.error(f"âš ï¸ Ù…Ø¬Ù„Ø¯ '{LAWS_DIR}/' ØºÙŠØ± Ù…ÙˆØ¬ÙˆØ¯. ÙŠØ±Ø¬Ù‰ Ø§Ù„ØªØ£ÙƒØ¯ Ù…Ù† ÙˆØ¬ÙˆØ¯ Ù…Ù„ÙØ§Øª Ø§Ù„Ù‚ÙˆØ§Ù†ÙŠÙ†.")
            return
        files = [f for f in os.listdir(LAWS_DIR) if f.endswith(".docx")]
        if not files:
            st.warning(f"ğŸ“‚ Ù„Ø§ ØªÙˆØ¬Ø¯ Ù…Ù„ÙØ§Øª Ù‚ÙˆØ§Ù†ÙŠÙ† ÙÙŠ Ù…Ø¬Ù„Ø¯ '{LAWS_DIR}/'.")
            return
        st.markdown("""
            <div style="direction: rtl; text-align: right;">
            <h3 style="display: flex; align-items: center; gap: 10px;">ğŸ” Ù†Ù…ÙˆØ°Ø¬ Ø§Ù„Ø¨Ø­Ø«</h3>
            </div>
        """, unsafe_allow_html=True)
        with st.form("main_search_form"):
            st.markdown('<div style="direction: rtl; text-align: right;">Ø§Ø®ØªØ± Ù‚Ø§Ù†ÙˆÙ†Ù‹Ø§ Ù„Ù„Ø¨Ø­Ø«:</div>', unsafe_allow_html=True)
            selected_file_form = st.selectbox("", ["Ø§Ù„ÙƒÙ„"] + files, key="main_file_select", label_visibility="collapsed")
            st.markdown('<div style="direction: rtl; text-align: right;">ğŸ“Œ Ø§ÙƒØªØ¨ ÙƒÙ„Ù…Ø© Ø£Ùˆ Ø¬Ù…Ù„Ø© Ù„Ù„Ø¨Ø­Ø« Ø¹Ù†Ù‡Ø§:</div>', unsafe_allow_html=True)
            st.markdown('<div dir="rtl">', unsafe_allow_html=True)
            keywords_form = st.text_area(
                "",
                key="main_keywords_input",
                help="Ø£Ø¯Ø®Ù„ Ø§Ù„ÙƒÙ„Ù…Ø§Øª Ø§Ù„ØªÙŠ ØªØ±ÙŠØ¯ Ø§Ù„Ø¨Ø­Ø« Ø¹Ù†Ù‡Ø§ØŒ ÙˆØ§ÙØµÙ„ Ø¨ÙŠÙ†Ù‡Ø§ Ø¨ÙØ§ØµÙ„Ø© Ø¥Ø°Ø§ ÙƒØ§Ù†Øª Ø£ÙƒØ«Ø± Ù…Ù† ÙƒÙ„Ù…Ø©.",
            )
            st.markdown('</div>', unsafe_allow_html=True)
            st.markdown('<div style="direction: rtl; text-align: right;">Ø£Ùˆ Ø£Ø¨Ø­Ø« Ø¨Ø±Ù‚Ù… Ø§Ù„Ù…Ø§Ø¯Ø©:</div>', unsafe_allow_html=True)
            st.markdown('<div dir="rtl">', unsafe_allow_html=True)
            article_number_input = st.text_input(
                "",
                key="article_number_input",
                help="Ø£Ø¯Ø®Ù„ Ø±Ù‚Ù… Ø§Ù„Ù…Ø§Ø¯Ø© Ù„Ù„Ø¨Ø­Ø« Ø¹Ù†Ù‡Ø§ Ù…Ø¨Ø§Ø´Ø±Ø© (ÙŠÙ…ÙƒÙ† Ø§Ø³ØªØ®Ø¯Ø§Ù… Ø£Ø±Ù‚Ø§Ù… Ø¹Ø±Ø¨ÙŠØ© Ø£Ùˆ Ø¥Ù†Ø¬Ù„ÙŠØ²ÙŠØ©)."
            )
            st.markdown('</div>', unsafe_allow_html=True)
            advanced_search_col = st.columns([1, 2, 5])
            with advanced_search_col[2]:
                exact_match = st.checkbox("ØªØ·Ø§Ø¨Ù‚ ØªØ§Ù… Ù„Ù„ÙƒÙ„Ù…Ø©", key="exact_match_checkbox")
            search_btn_col = st.columns([1, 2, 12])
            with search_btn_col[2]:
                submitted = st.form_submit_button("ğŸ” Ø¨Ø¯Ø¡ Ø§Ù„Ø¨Ø­Ø«", use_container_width=True)
        if "results" not in st.session_state:
            st.session_state.results = []
        if "search_done" not in st.session_state:
            st.session_state.search_done = False
        if submitted:
            results = []
            search_files = files if selected_file_form == "Ø§Ù„ÙƒÙ„" else [selected_file_form]
            kw_list = [k.strip() for k in keywords_form.split(",") if k.strip()] if keywords_form else []
            search_by_article = bool(article_number_input.strip())
            normalized_kw_list = [normalize_arabic_text(kw) for kw in kw_list] if kw_list else []
            norm_article = normalize_arabic_numbers(article_number_input.strip()) if search_by_article else ""
            with st.spinner("Ø¬Ø§Ø±ÙŠ Ø§Ù„Ø¨Ø­Ø« ÙÙŠ Ø§Ù„Ù‚ÙˆØ§Ù†ÙŠÙ†... Ù‚Ø¯ ÙŠØ³ØªØºØ±Ù‚ Ø§Ù„Ø£Ù…Ø± Ø¨Ø¹Ø¶ Ø§Ù„ÙˆÙ‚Øª."):
                for file in search_files:
                    try:
                        doc = Document(os.path.join(LAWS_DIR, file))
                    except Exception as e:
                        st.warning(f"âš ï¸ ØªØ¹Ø°Ø± Ù‚Ø±Ø§Ø¡Ø© Ø§Ù„Ù…Ù„Ù {file}: {e}. ÙŠØ±Ø¬Ù‰ Ø§Ù„ØªØ£ÙƒØ¯ Ù…Ù† Ø£Ù†Ù‡ Ù…Ù„Ù DOCX ØµØ§Ù„Ø­.")
                        continue
                    law_name = file.replace(".docx", "")
                    last_article = "ØºÙŠØ± Ù…Ø¹Ø±ÙˆÙØ©"
                    current_article_paragraphs = []
                    for para in doc.paragraphs:
                        txt = para.text.strip()
                        if not txt:
                            continue
                        match = re.match(r"Ù…Ø§Ø¯Ø©\s*[\(]?\s*(\d+)[\)]?", txt)
                        if match:
                            if current_article_paragraphs:
                                full_text = "\n".join(current_article_paragraphs)
                                add_result = False
                                simple_full_text = normalize_arabic_text(full_text)
                                if search_by_article and normalize_arabic_numbers(last_article) == norm_article:
                                    add_result = True
                                elif normalized_kw_list:
                                    for idx, kw in enumerate(normalized_kw_list):
                                        if not kw:
                                            continue
                                        if exact_match:
                                            pattern = r'(?<!\w)'+re.escape(kw)+r'(?!\w)'
                                            if re.search(pattern, simple_full_text):
                                                add_result = True
                                                break
                                        else:
                                            if kw in simple_full_text:
                                                add_result = True
                                                break
                                if add_result:
                                    highlighted = highlight_keywords(full_text, kw_list, normalized_keywords=normalized_kw_list, exact_match=exact_match) if kw_list else full_text
                                    results.append({
                                        "law": law_name,
                                        "num": last_article,
                                        "text": highlighted,
                                        "plain": full_text
                                    })
                                current_article_paragraphs = []
                            last_article = match.group(1)
                        current_article_paragraphs.append(txt)
                    if current_article_paragraphs:
                        full_text = "\n".join(current_article_paragraphs)
                        add_result = False
                        simple_full_text = normalize_arabic_text(full_text)
                        if search_by_article and normalize_arabic_numbers(last_article) == norm_article:
                            add_result = True
                        elif normalized_kw_list:
                            for idx, kw in enumerate(normalized_kw_list):
                                if not kw:
                                    continue
                                if exact_match:
                                    pattern = r'(?<!\w)'+re.escape(kw)+r'(?!\w)'
                                    if re.search(pattern, simple_full_text):
                                        add_result = True
                                        break
                                else:
                                    if kw in simple_full_text:
                                        add_result = True
                                        break
                        if add_result:
                            highlighted = highlight_keywords(full_text, kw_list, normalized_keywords=normalized_kw_list, exact_match=exact_match) if kw_list else full_text
                            results.append({
                                "law": law_name,
                                "num": last_article,
                                "text": highlighted,
                                "plain": full_text
                            })
            st.session_state.results = results
            st.session_state.search_done = True
            if not results:
                st.info("Ù„Ù… ÙŠØªÙ… Ø§Ù„Ø¹Ø«ÙˆØ± Ø¹Ù„Ù‰ Ù†ØªØ§Ø¦Ø¬ Ù…Ø·Ø§Ø¨Ù‚Ø© Ù„Ù„Ø¨Ø­Ø«.")
        if st.session_state.get("search_done", False) and st.session_state.results:
            st.markdown("<h2 style='text-align: center; color: #388E3C;'>Ù†ØªØ§Ø¦Ø¬ Ø§Ù„Ø¨Ø­Ø« ÙÙŠ Ø§Ù„Ù‚ÙˆØ§Ù†ÙŠÙ† ğŸ“š</h2>", unsafe_allow_html=True)
            st.markdown("---")
        if st.session_state.get("search_done", False):
            results = st.session_state.results
            unique_laws = sorted(set(r["law"] for r in results))
            st.markdown('<div class="rtl-metric">', unsafe_allow_html=True)
            st.metric(label="ğŸ“Š Ø¥Ø¬Ù…Ø§Ù„ÙŠ Ø§Ù„Ù†ØªØ§Ø¦Ø¬ Ø§Ù„ØªÙŠ ØªÙ… Ø§Ù„Ø¹Ø«ÙˆØ± Ø¹Ù„ÙŠÙ‡Ø§", value=f"{len(results)}", delta=f"ÙÙŠ {len(unique_laws)} Ù‚Ø§Ù†ÙˆÙ†/Ù…Ù„Ù")
            st.markdown('</div>', unsafe_allow_html=True)
            if results:
                export_data = export_results_to_word(results)
                st.markdown('<div class="rtl-download-btn">', unsafe_allow_html=True)
                st.download_button(
                    label="â¬‡ï¸ ØªØµØ¯ÙŠØ± Ø§Ù„Ù†ØªØ§Ø¦Ø¬ Ø¥Ù„Ù‰ Word",
                    data=export_data,
                    file_name="Ù†ØªØ§Ø¦Ø¬_Ø§Ù„Ø¨Ø­Ø«_Ø§Ù„Ù‚ÙˆØ§Ù†ÙŠÙ†_Ø§Ù„ÙŠÙ…Ù†ÙŠØ©.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_button_word_main",
                    use_container_width=False
                )
                st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.warning("Ù„Ø§ ØªÙˆØ¬Ø¯ Ù†ØªØ§Ø¦Ø¬ Ù„ØªØµØ¯ÙŠØ±Ù‡Ø§.")
            st.markdown("---")
            if results:
                for i, r in enumerate(results):
                    with st.expander(f"ğŸ“š Ø§Ù„Ù…Ø§Ø¯Ø© ({r['num']}) Ù…Ù† Ù‚Ø§Ù†ÙˆÙ† {r['law']}", expanded=True):
                        st.markdown(f'''
                        <div class="result-box-night">
                            <p style="font-size:17px;line-height:1.8;margin-top:0px;">
                                {r["text"]}
                            </p>
                        </div>
                        ''', unsafe_allow_html=True)
                        components.html(f"""
                            <style>
                            .copy-material-btn {{
                                display: inline-flex;
                                align-items: center;
                                gap: 10px;
                                background: linear-gradient(90deg, #1abc9c 0%, #2980b9 100%);
                                color: #fff;
                                border: none;
                                border-radius: 30px;
                                font-size: 18px;
                                font-family: 'Cairo', 'Tajawal', sans-serif;
                                padding: 10px 22px;
                                cursor: pointer;
                                box-shadow: 0 4px 15px rgba(41, 128, 185, 0.4);
                                transition: all 0.3s ease;
                                margin-bottom: 10px;
                                direction: rtl;
                                white-space: nowrap;
                            }}
                            .copy-material-btn:hover {{
                                background: linear-gradient(90deg, #2980b9 0%, #1abc9c 100%);
                                box-shadow: 0 6px 20px rgba(41, 128, 185, 0.6);
                                transform: translateY(-2px);
                            }}
                            .copy-material-btn .copy-icon {{
                                font-size: 20px;
                                margin-left: 8px;
                                display: block;
                            }}
                            .copy-material-btn .copied-check {{
                                font-size: 20px;
                                color: #fff;
                                margin-left: 8px;
                                display: none;
                            }}
                            .copy-material-btn.copied .copy-icon {{
                                display: none;
                            }}
                            .copy-material-btn.copied .copied-check {{
                                display: inline;
                                animation: fadein-check 0.5s ease-out;
                            }}
                            @keyframes fadein-check {{
                                0% {{ opacity: 0; transform: scale(0.7); }}
                                100% {{ opacity: 1; transform: scale(1); }}
                            }}
                            </style>
                            <button class="copy-material-btn" id="copy_btn_{i}_{r['law']}_{r['num']}" onclick="
                                navigator.clipboard.writeText(document.getElementById('plain_text_{i}_{r['law']}_{r['num']}').innerText);
                                var btn = document.getElementById('copy_btn_{i}_{r['law']}_{r['num']}');
                                btn.classList.add('copied');
                                setTimeout(function(){{
                                    btn.classList.remove('copied');
                                }}, 1800);
                            ">
                                <span class="copy-icon">
                                    <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
                                        <rect x="9" y="9" width="13" height="13" rx="2" ry="2"></rect>
                                        <path d="M5 15H4a2 2 0 0 1-2-2V4a2 2 0 0 1 2-2h9a2 2 0 0 1 2 2v1"></path>
                                    </svg>
                                </span>
                                <span>Ù†Ø³Ø®</span>
                                <span class="copied-check">
                                    <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
                                        <polyline points="20 6 9 17 4 12"></polyline>
                                    </svg>
                                    ØªÙ… Ø§Ù„Ù†Ø³Ø®!
                                </span>
                            </button>
                            <div id="plain_text_{i}_{r['law']}_{r['num']}" style="display:none;">{html.escape(r['plain'])}</div>
                        """, height=60)
            else:
                st.info("Ù„Ø§ ØªÙˆØ¬Ø¯ Ù†ØªØ§Ø¦Ø¬ Ù„Ø¹Ø±Ø¶Ù‡Ø§ Ø­Ø§Ù„ÙŠÙ‹Ø§. ÙŠØ±Ø¬Ù‰ Ø¥Ø¬Ø±Ø§Ø¡ Ø¨Ø­Ø« Ø¬Ø¯ÙŠØ¯.")
    with tabs[1]:
        if not os.path.exists(LAWS_DIR):
            st.error(f"âš ï¸ Ù…Ø¬Ù„Ø¯ '{LAWS_DIR}/' ØºÙŠØ± Ù…ÙˆØ¬ÙˆØ¯. ÙŠØ±Ø¬Ù‰ Ø§Ù„ØªØ£ÙƒØ¯ Ù…Ù† ÙˆØ¬ÙˆØ¯ Ù…Ù„ÙØ§Øª Ø§Ù„Ù‚ÙˆØ§Ù†ÙŠÙ†.")
            return
        files = [f for f in os.listdir(LAWS_DIR) if f.endswith(".docx")]
        if not files:
            st.warning(f"ğŸ“‚ Ù„Ø§ ØªÙˆØ¬Ø¯ Ù…Ù„ÙØ§Øª Ù‚ÙˆØ§Ù†ÙŠÙ† ÙÙŠ Ù…Ø¬Ù„Ø¯ '{LAWS_DIR}/'.")
            return
        render_law_file_viewer(files)

def render_header():
    if os.path.exists("header.html"):
        with open("header.html", "r", encoding="utf-8") as f:
            header_html = f.read()
        st.markdown(header_html, unsafe_allow_html=True)
    else:
        st.error("âš ï¸ Ù…Ù„Ù 'header.html' ØºÙŠØ± Ù…ÙˆØ¬ÙˆØ¯ ÙÙŠ Ù…Ø¬Ù„Ø¯ Ø§Ù„Ù…Ø´Ø±ÙˆØ¹.")

def main():
    render_header()
    device_id = get_device_id()
    trial_start = get_trial_start(device_id)
    if is_activated():
        run_main_app()
        return
    if trial_start is not None:
        elapsed_time = time.time() - trial_start
        remaining_time = int(TRIAL_DURATION - elapsed_time)
        if remaining_time > 0:
            run_main_app()
            return
        else:
            st.error("âŒ Ø§Ù†ØªÙ‡Øª Ù…Ø¯Ø© Ø§Ù„ØªØ¬Ø±Ø¨Ø© Ø§Ù„Ù…Ø¬Ø§Ù†ÙŠØ© Ù„Ù‡Ø°Ø§ Ø§Ù„Ø¬Ù‡Ø§Ø². ÙŠØ±Ø¬Ù‰ ØªÙØ¹ÙŠÙ„ Ø§Ù„ØªØ·Ø¨ÙŠÙ‚ Ù„Ù„Ø§Ø³ØªÙ…Ø±Ø§Ø± ÙÙŠ Ø§Ù„Ø§Ø³ØªØ®Ø¯Ø§Ù….")
    with st.container(border=True):
        if trial_start is None:
            if st.button("ğŸš€ Ø¨Ø¯Ø¡ Ø§Ù„Ù†Ø³Ø®Ø© Ø§Ù„Ù…Ø¬Ø§Ù†ÙŠØ©", key="start_trial_button", use_container_width=True):
                register_trial(device_id)
                st.rerun()
    st.markdown("---")
    with st.container(border=True):
        st.markdown("<h3 style='text-align:center; color:#2c3e50;'>ğŸ” Ø§Ù„Ù†Ø³Ø®Ø© Ø§Ù„Ù…Ø¯ÙÙˆØ¹Ø©</h3>", unsafe_allow_html=True)
        code = st.text_input("Ø£Ø¯Ø®Ù„ ÙƒÙˆØ¯ Ø§Ù„ØªÙØ¹ÙŠÙ„ Ù‡Ù†Ø§:", key="activation_code_input", help="Ø£Ø¯Ø®Ù„ Ø§Ù„ÙƒÙˆØ¯ Ø§Ù„Ø°ÙŠ Ø­ØµÙ„Øª Ø¹Ù„ÙŠÙ‡ Ù„ØªÙØ¹ÙŠÙ„ Ø§Ù„Ù†Ø³Ø®Ø© Ø§Ù„ÙƒØ§Ù…Ù„Ø©.")
        if st.button("âœ… ØªÙØ¹ÙŠÙ„ Ø§Ù„Ø¢Ù†", key="activate_button", use_container_width=True):
            if code and activate_app(code.strip()):
                st.success("âœ… ØªÙ… Ø§Ù„ØªÙØ¹ÙŠÙ„ Ø¨Ù†Ø¬Ø§Ø­! ÙŠØ±Ø¬Ù‰ Ø¥Ø¹Ø§Ø¯Ø© ØªØ´ØºÙŠÙ„ Ø§Ù„ØªØ·Ø¨ÙŠÙ‚ Ù„ØªØ·Ø¨ÙŠÙ‚ Ø§Ù„ØªØºÙŠÙŠØ±Ø§Øª.")
                st.stop()
            else:
                st.error("âŒ ÙƒÙˆØ¯ Ø§Ù„ØªÙØ¹ÙŠÙ„ ØºÙŠØ± ØµØ­ÙŠØ­ Ø£Ùˆ Ø§Ù†ØªÙ‡Øª ØµÙ„Ø§Ø­ÙŠØªÙ‡.")

if __name__ == "__main__":
    main()